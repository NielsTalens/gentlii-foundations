from pathlib import Path

from docx import Document

from gentlii_foundations.extractors.docx import extract_docx_text


def test_extract_docx_text_reads_paragraphs(tmp_path: Path):
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First line")
    doc.add_paragraph("Second line")
    doc.save(path)

    result = extract_docx_text(path)

    assert "First line" in result
    assert "Second line" in result


def test_extract_docx_text_reads_table_cell_text_from_fixture():
    path = Path(__file__).resolve().parent.parent / "fixtures" / "docx" / "table-only-content.docx"

    result = extract_docx_text(path)

    assert "Critical requirement lives in a table cell" in result
    assert "Launch blocker hidden in table" in result
