from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.table import Table


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = _extract_table_text(document.tables)
    return "\n\n".join(paragraphs + table_cells)


def _extract_table_text(tables: list[Table]) -> list[str]:
    text: list[str] = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                text.extend(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
                text.extend(_extract_table_text(cell.tables))
    return text
